# -*- coding: utf-8 -*-
"""舊版 .doc handler：LibreOffice 轉 docx → 沿用 docx handler 取代 → 轉回 .doc。

R2：進來是 .doc，出去也一定是 .doc，不做「順手升級」。
"""
from __future__ import annotations

from typing import Iterable

from docx import Document

from src import legacy_office
from src.config import MatchOptions
from src.handlers import docx_file

EXTENSIONS = (".doc",)


def _fingerprint(docx_path: str):
    """內容量指紋：段落數 + 表格數 + 非空段落的文字總長度。"""
    document = Document(docx_path)
    tables = len(document.tables)
    paragraphs = [p.text for p in document.paragraphs]
    return (len(paragraphs), tables, sum(len(t) for t in paragraphs))


def process(src_path: str, dst_path: str, keywords: Iterable[str], options: MatchOptions):
    keywords = list(keywords)

    def replace(modern_in, modern_out):
        outcome = docx_file.process(modern_in, modern_out, keywords, options)
        return outcome if isinstance(outcome, tuple) else (outcome, "")

    count, message = legacy_office.round_trip(
        src_path, "docx", "doc", dst_path, replace, _fingerprint)
    notes = [n for n in (message, "經 LibreOffice 轉檔，複雜排版可能有差異") if n]
    return count, "；".join(notes)
