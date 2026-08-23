# -*- coding: utf-8 -*-
"""重新產生 docx / xlsx 測試樣本（樣本已入庫，只有要調整內容時才需重跑）。

    python3 tests/make_samples.py
"""
from __future__ import annotations

import os

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font
from openpyxl.utils import get_column_letter

SAMPLES = os.path.join(os.path.dirname(os.path.abspath(__file__)), "samples")


def make_docx():
    doc = Document()
    # 關鍵字被切成三個 run，模擬 Word 實際的 run 切法
    p = doc.add_paragraph()
    for text in ("這是", "機", "密", "文件"):
        p.add_run(text)

    # 有樣式的 run：取代後樣式必須留著
    p2 = doc.add_paragraph()
    run = p2.add_run("機密報告")
    run.bold = True
    run.font.size = Pt(16)

    # 表格
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "機密"
    table.cell(0, 1).text = "一般"

    # 頁首頁尾
    section = doc.sections[0]
    section.header.paragraphs[0].text = "頁首：機密"
    section.footer.paragraphs[0].text = "頁尾：無關鍵字"

    doc.save(os.path.join(SAMPLES, "split_runs.docx"))


def make_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "資料"
    ws["A1"] = "機密資料"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = "一般資料"
    ws["B1"] = '=CONCATENATE("機密-",A1)'   # 字串常數含關鍵字 → 要換
    ws["B2"] = "=SUM(C1:C9)"                 # 純參照 → 不得動
    ws["C1"] = 123                            # 數值 → 不得動
    ws.merge_cells("A4:B4")
    ws["A4"] = "合併儲存格內的機密"
    ws.column_dimensions[get_column_letter(1)].width = 28

    ws2 = wb.create_sheet("第二頁")
    ws2["A1"] = "機密"
    wb.save(os.path.join(SAMPLES, "book.xlsx"))


def make_pdfs():
    import fitz

    # 文字層 PDF：中文 + 英文 + 一張圖，另加一頁不含關鍵字
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 100), "這是機密文件", fontname="china-t", fontsize=14)
    page.insert_text((72, 140), "Project Secret-01 是機密的", fontname="china-t", fontsize=12)
    page.insert_text((72, 180), "紅字機密", fontname="china-t", fontsize=18, color=(0.8, 0, 0))
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 60, 40))
    pix.set_rect(pix.irect, (30, 120, 200))
    page.insert_image(fitz.Rect(400, 90, 460, 130), pixmap=pix)
    page2 = doc.new_page()
    page2.insert_text((72, 100), "第二頁沒有關鍵字", fontname="china-t", fontsize=12)
    doc.save(os.path.join(SAMPLES, "text_layer.pdf"))
    doc.close()

    # 掃描影像 PDF：整頁只有一張圖，沒有任何文字層
    doc = fitz.open()
    page = doc.new_page()
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 400, 200))
    pix.set_rect(pix.irect, (220, 220, 220))
    page.insert_image(fitz.Rect(72, 72, 472, 272), pixmap=pix)
    doc.save(os.path.join(SAMPLES, "scanned.pdf"))
    doc.close()


def _text_image(lines, size=(560, 220)):
    """畫一張含文字的圖當 OCR 樣本；找得到 CJK 字型就用，找不到就退回英文字型。"""
    from PIL import Image, ImageDraw, ImageFont

    cjk_candidates = (
        ("/System/Library/Fonts/PingFang.ttc", 0),
        ("/System/Library/Fonts/STHeiti Medium.ttc", 0),
        ("C:/Windows/Fonts/msjh.ttc", 0),
        ("/System/Library/Fonts/Supplemental/Songti.ttc", 0),
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", None),
    )

    def _renders_all(candidate, sample):
        """Songti.ttc 少了「這」「機」等字形，會畫成空白——必須逐字驗過才用。"""
        for char in sample:
            probe = Image.new("L", (56, 64), 255)
            ImageDraw.Draw(probe).text((2, 2), char, fill=0, font=candidate)
            if not any(pixel < 128 for pixel in probe.getdata()):
                return False
        return True

    font = None
    sample_text = [c for c in "".join(lines) if c.strip()]
    for path, index in cjk_candidates:
        if not os.path.exists(path):
            continue
        try:
            candidate = (ImageFont.truetype(path, 40, index=index) if index is not None
                         else ImageFont.truetype(path, 40))
        except OSError:
            continue
        if _renders_all(candidate, sample_text):
            font = candidate
            break
    if font is None:
        font = ImageFont.load_default()

    image = Image.new("RGB", size, (255, 255, 255))
    draw = ImageDraw.Draw(image)
    y = 24
    for line in lines:
        draw.text((24, y), line, fill=(0, 0, 0), font=font)
        y += 60
    return image


def make_image_samples():
    """OCR 樣本：英文行讓只裝 eng 的環境也測得到，中文行測 chi_tra。"""
    img = _text_image(["Project Secret 01", "這是機密文件"])
    img.save(os.path.join(SAMPLES, "ocr_page.png"))

    # 掃描影像 PDF：把同一張圖鋪成整頁，沒有任何文字層
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    buf = os.path.join(SAMPLES, "ocr_page.png")
    page.insert_image(fitz.Rect(50, 50, 545, 245), filename=buf)
    doc.save(os.path.join(SAMPLES, "scanned_text.pdf"))
    doc.close()

    # JPEG 樣本（含 EXIF，用來驗證重新編碼後有沒有保住）
    jpeg = _text_image(["Project Secret 01", "這是機密文件"]).convert("RGB")
    exif = jpeg.getexif()
    exif[270] = "snoopy test sample"          # ImageDescription
    exif[305] = "make_samples.py"             # Software
    jpeg.save(os.path.join(SAMPLES, "ocr_page.jpg"), format="JPEG",
              quality=92, exif=exif.tobytes())

    # docx 內嵌圖片
    from docx import Document as Doc
    d = Doc()
    d.add_paragraph("文件本文沒有關鍵字")
    d.add_picture(buf)
    d.save(os.path.join(SAMPLES, "with_image.docx"))


def make_legacy_samples():
    """用 LibreOffice 把 docx / xlsx 樣本轉成舊版 .doc / .xls；沒裝就跳過。"""
    import sys

    sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    from src import legacy_office

    if not legacy_office.is_available():
        print("略過 .doc / .xls 樣本：找不到 LibreOffice")
        return

    for source, target in (("split_runs.docx", "doc"), ("book.xlsx", "xls")):
        produced = legacy_office.convert(os.path.join(SAMPLES, source), target, SAMPLES)
        print("已產生", os.path.basename(produced))


if __name__ == "__main__":
    make_docx()
    make_xlsx()
    make_pdfs()
    make_image_samples()
    make_legacy_samples()
    print("樣本已重新產生於 %s" % SAMPLES)
