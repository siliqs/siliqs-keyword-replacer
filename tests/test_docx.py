# -*- coding: utf-8 -*-
import io
import os

from docx import Document

from conftest import SAMPLES, call
from src.config import MatchOptions
from src.handlers import docx_file

SRC = os.path.join(SAMPLES, "split_runs.docx")


def _process(tmp_path, keywords=("機密",), options=None):
    dst = str(tmp_path / "out.docx")
    count, _ = call(docx_file, SRC, dst, keywords, options or MatchOptions())
    return Document(dst), count, dst


def test_replaces_keyword_split_across_runs(tmp_path):
    doc, count, _ = _process(tmp_path)
    assert doc.paragraphs[0].text == "這是snoopy文件"     # 原本被切成「機」「密」兩個 run
    assert count == 4                                      # 拆散的1 + 有樣式的1 + 表格1 + 頁首1


def test_keeps_run_formatting(tmp_path):
    doc, _, _ = _process(tmp_path)
    styled = [p for p in doc.paragraphs if p.text == "snoopy報告"][0]
    run = styled.runs[0]
    assert run.bold is True
    assert run.font.size.pt == 16.0                        # R5：只有文字變，樣式不動


def test_replaces_in_tables_and_header(tmp_path):
    doc, _, _ = _process(tmp_path)
    table = doc.tables[0]
    assert table.cell(0, 0).text == "snoopy"
    assert table.cell(0, 1).text == "一般"
    header = doc.sections[0].header.paragraphs[0].text
    footer = doc.sections[0].footer.paragraphs[0].text
    assert header == "頁首：snoopy"
    assert footer == "頁尾：無關鍵字"


def test_source_untouched(tmp_path):
    before = io.open(SRC, "rb").read()
    _process(tmp_path)
    assert io.open(SRC, "rb").read() == before             # R1


def test_no_keyword_no_change(tmp_path):
    doc, count, _ = _process(tmp_path, keywords=("不存在的詞",))
    assert count == 0
    assert doc.paragraphs[0].text == "這是機密文件"
